"""
将 Part1_基础篇.docx 转换为多篇微信公众号 HTML 文章
- 图片内嵌 base64（可直接复制粘贴到微信编辑器）
- 自动压缩图片控制体积
- 正确处理 Word 表格
- 微信公众号优化排版
"""
import sys
import base64
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image

# ============================================================
# 文章拆分方案（7篇），按段落索引切分
# ============================================================
ARTICLES = [
    {
        "file": "01_走进AI大模型的世界.html",
        "title": "AI大模型到底是什么？一篇文章带你从零搞懂",
        "subtitle": "从猜水果到GPT-4，人工智能的三次浪潮与大模型的前世今生",
        "start": 2, "end": 103,
    },
    {
        "file": "02_第一次调用大模型API.html",
        "title": "5分钟，让你亲手调用一次大模型API",
        "subtitle": "零代码经验也能跑通的AI实战第一课",
        "start": 103, "end": 194,
    },
    {
        "file": "03_Python环境与基础语法.html",
        "title": "AI开发第一步：Python环境搭建与语法速成",
        "subtitle": "Conda环境、变量类型、字符串操作、条件循环，一文搞定",
        "start": 194, "end": 423,
    },
    {
        "file": "04_Python进阶与数据处理.html",
        "title": "Python进阶：函数、类、NumPy、Pandas一网打尽",
        "subtitle": "从模块化编程到数据处理，AI开发者的必备技能",
        "start": 423, "end": 865,
    },
    {
        "file": "05_深度学习核心概念.html",
        "title": "深度学习没那么难：从神经网络到PyTorch实战",
        "subtitle": "感知机、损失函数、反向传播、CNN/RNN、GPU配置与手写数字识别",
        "start": 865, "end": 1180,
    },
    {
        "file": "06_NLP基础与词嵌入.html",
        "title": "让机器读懂人话：NLP基础与词嵌入技术详解",
        "subtitle": "分词、TF-IDF、Word2Vec、语言模型演进，一文讲透",
        "start": 1180, "end": 1326,
    },
    {
        "file": "07_HuggingFace与情感分析实战.html",
        "title": "手把手教你用HuggingFace做中文情感分析",
        "subtitle": "从Transformers库入门到完整实战项目，附完整代码",
        "start": 1326, "end": 99999,
    },
]

# XML namespaces
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
WP_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

IMG_MAX_WIDTH = 1080
IMG_JPEG_QUALITY = 82

# ============================================================
# 微信公众号 CSS
# ============================================================
WECHAT_CSS = """\
<style>
* { margin: 0; padding: 0; box-sizing: border-box; }
body {
  font-family: -apple-system, BlinkMacSystemFont, "Helvetica Neue",
               "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", sans-serif;
  font-size: 15px; line-height: 2; color: #3f3f3f; background: #fff;
  letter-spacing: 0.5px;
}
</style>"""

# ============================================================
# Inline styles — 微信粘贴时 <style> 块会被丢弃，必须用 inline style
# ============================================================
S_WRAP = 'max-width:100%;padding:24px 16px;'
S_TITLE = 'font-size:24px;font-weight:bold;color:#1a1a1a;text-align:center;margin-bottom:6px;line-height:1.5;'
S_SUB = 'font-size:13px;color:#999;text-align:center;margin-bottom:28px;padding-bottom:20px;border-bottom:1px solid #eee;line-height:1.6;'
S_P = 'margin:0 0 18px 0;text-align:justify;text-indent:2em;font-size:15px;line-height:2;color:#3f3f3f;'
S_P_NI = 'margin:0 0 18px 0;text-align:justify;font-size:15px;line-height:2;color:#3f3f3f;'
S_H1 = 'font-size:20px;font-weight:bold;color:#fff;text-align:center;margin:36px 0 24px;padding:10px 20px;background:linear-gradient(135deg,#07c160,#06ad56);border-radius:6px;letter-spacing:1px;'
S_H2 = 'font-size:18px;font-weight:bold;color:#07c160;margin:30px 0 16px;padding:6px 0 6px 14px;border-left:4px solid #07c160;background:linear-gradient(90deg,#f0faf4,transparent);'
S_H3 = 'font-size:16px;font-weight:bold;color:#333;margin:24px 0 12px;padding-left:12px;border-left:3px solid #ccc;'
S_H4 = 'font-size:15px;font-weight:bold;color:#555;margin:20px 0 10px;'
S_NOTE = 'background:#f6ffed;border-left:4px solid #52c41a;padding:14px 18px;margin:20px 0;border-radius:0 8px 8px 0;font-size:14px;color:#666;line-height:1.8;'
S_CODE_WRAP = 'background:#2b2b2b;border-radius:8px;padding:16px 18px;margin:18px 0;overflow-x:auto;-webkit-overflow-scrolling:touch;'
S_CODE_LINE = "font-family:Menlo,Consolas,'Courier New',monospace;font-size:13px;line-height:1.7;color:#e6e6e6;margin:0;padding:0;border:none;background:none;"
S_IW = 'text-align:center;margin:22px 0;'
S_IMG = 'max-width:100%;height:auto;border-radius:6px;box-shadow:0 4px 12px rgba(0,0,0,0.08);'
S_IC = 'font-size:12px;color:#aaa;text-align:center;margin-top:8px;line-height:1.6;'
S_LI = 'margin-bottom:10px;line-height:1.8;font-size:15px;color:#3f3f3f;'
S_TH = 'background:#07c160;color:#fff;font-weight:bold;padding:10px 12px;text-align:left;border:1px solid #07c160;font-size:13px;'
S_TD = 'padding:10px 12px;text-align:left;border:1px solid #e8e8e8;color:#3f3f3f;font-size:13px;line-height:1.6;'
S_TD_EVEN = S_TD + 'background:#f7faf8;'
S_SEP = 'text-align:center;margin:28px 0;color:#ddd;font-size:20px;letter-spacing:10px;'
S_FOOTER = 'margin-top:36px;padding-top:20px;border-top:1px solid #eee;font-size:13px;color:#bbb;text-align:center;line-height:1.8;'
S_STRONG = 'color:#07c160;font-weight:bold;'
S_EM = 'font-style:normal;color:#d4380d;border-bottom:1px dashed #d4380d;'


def escape(text: str) -> str:
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def escape_code_line(text: str) -> str:
    """Escape a code line and convert leading spaces to &nbsp; for WeChat."""
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    # Convert ALL spaces to &nbsp; so indentation survives WeChat paste
    text = text.replace(' ', '&nbsp;')
    if not text:
        text = '&nbsp;'  # empty lines need content to keep height
    return text


def compress_image_to_base64(blob: bytes, content_type: str) -> str:
    """Compress image and return base64 data URI."""
    try:
        img = Image.open(BytesIO(blob))
    except Exception:
        b64 = base64.b64encode(blob).decode('ascii')
        return f"data:{content_type};base64,{b64}"

    if img.width > IMG_MAX_WIDTH:
        ratio = IMG_MAX_WIDTH / img.width
        img = img.resize((IMG_MAX_WIDTH, int(img.height * ratio)), Image.LANCZOS)

    has_real_alpha = False
    if img.mode in ('RGBA', 'LA', 'PA') and img.mode == 'RGBA':
        if img.getchannel('A').getextrema()[0] < 250:
            has_real_alpha = True

    buf = BytesIO()
    if has_real_alpha:
        img.save(buf, format='PNG', optimize=True)
        mime = 'image/png'
    else:
        img.convert('RGB').save(buf, format='JPEG', quality=IMG_JPEG_QUALITY, optimize=True)
        mime = 'image/jpeg'

    b64 = base64.b64encode(buf.getvalue()).decode('ascii')
    return f"data:{mime};base64,{b64}"


def extract_image_base64(para_element, doc_part) -> str | None:
    """Extract image from paragraph XML element, return base64 data URI."""
    blips = para_element.findall(f'.//{{{A_NS}}}blip')
    if not blips:
        return None
    embed_id = blips[0].get(f'{{{R_NS}}}embed')
    if not embed_id:
        return None
    try:
        image_part = doc_part.related_parts[embed_id]
        return compress_image_to_base64(image_part.blob, image_part.content_type)
    except (KeyError, AttributeError):
        return None


def get_para_text(elem) -> str:
    """Get full text from a <w:p> element."""
    texts = []
    for t in elem.iter(qn('w:t')):
        if t.text:
            texts.append(t.text)
    return ''.join(texts)


def get_para_style(elem) -> str:
    """Get style name from <w:p> element."""
    pPr = elem.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            return pStyle.get(qn('w:val'), 'Normal')
    return 'Normal'


def build_rich_text_from_element(elem) -> str:
    """Build HTML with bold/italic from <w:p> element runs."""
    parts = []
    for run_elem in elem.findall(qn('w:r')):
        t_elems = run_elem.findall(qn('w:t'))
        t = ''.join((te.text or '') for te in t_elems)
        if not t:
            continue
        t = escape(t)

        rPr = run_elem.find(qn('w:rPr'))
        is_bold = False
        is_italic = False
        if rPr is not None:
            b = rPr.find(qn('w:b'))
            if b is not None and b.get(qn('w:val'), 'true') != 'false':
                is_bold = True
            i = rPr.find(qn('w:i'))
            if i is not None and i.get(qn('w:val'), 'true') != 'false':
                is_italic = True

        if is_bold and is_italic:
            t = f'<strong style="{S_STRONG}"><em style="{S_EM}">{t}</em></strong>'
        elif is_bold:
            t = f'<strong style="{S_STRONG}">{t}</strong>'
        elif is_italic:
            t = f'<em style="{S_EM}">{t}</em>'
        parts.append(t)

    result = ''.join(parts)
    if not result:
        result = escape(get_para_text(elem))
    return result


def table_to_html(tbl_element, doc) -> str:
    """Convert a <w:tbl> element to HTML table."""
    table = Table(tbl_element, doc)
    rows = table.rows
    if not rows:
        return ''

    html = ['<section style="overflow-x:auto;margin:20px 0;-webkit-overflow-scrolling:touch;"><table style="width:100%;border-collapse:collapse;font-size:13px;">']

    for r_idx, row in enumerate(rows):
        html.append('<tr>')
        for cell in row.cells:
            if r_idx == 0:
                tag, sty = 'th', S_TH
            else:
                sty = S_TD_EVEN if r_idx % 2 == 0 else S_TD
                tag = 'td'
            cell_texts = []
            for p in cell.paragraphs:
                t = p.text.strip()
                if t:
                    cell_texts.append(escape(t))
            cell_html = '<br>'.join(cell_texts) if cell_texts else '&nbsp;'
            html.append(f'<{tag} style="{sty}">{cell_html}</{tag}>')
        html.append('</tr>')

    html.append('</table></section>')
    return '\n'.join(html)


def build_body_elements(doc):
    """Build ordered list of (type, element) from document body.
    type is 'p' for paragraph, 'tbl' for table.
    For paragraphs, also returns the para index (matching doc.paragraphs)."""
    body = doc.element.body
    items = []
    para_idx = 0
    for child in body:
        tag = etree.QName(child).localname
        if tag == 'p':
            items.append(('p', child, para_idx))
            para_idx += 1
        elif tag == 'tbl':
            items.append(('tbl', child, para_idx))
        # skip other elements (sectPr etc.)
    return items


def convert_chapter(doc, start_idx: int, end_idx: int,
                    title: str, subtitle: str) -> str:
    """Convert a paragraph range to WeChat HTML, including tables."""
    doc_part = doc.part
    elements = build_body_elements(doc)
    end_idx = min(end_idx, len(doc.paragraphs))

    o = []
    o.append(f'<section style="{S_WRAP}">')
    o.append(f'<section style="{S_TITLE}">{escape(title)}</section>')
    o.append(f'<section style="{S_SUB}">{escape(subtitle)}</section>')

    i = 0
    total = len(elements)
    while i < total:
        etype, elem, pidx = elements[i]

        # --- Table ---
        if etype == 'tbl':
            if start_idx <= pidx < end_idx:
                o.append(table_to_html(elem, doc))
            i += 1
            continue

        # --- Paragraph: check range ---
        if pidx < start_idx:
            i += 1
            continue
        if pidx >= end_idx:
            break

        style = get_para_style(elem)
        text = get_para_text(elem).strip()

        # Code blocks: batch consecutive, each line a <p> with &nbsp;
        if style == '代码清单':
            code_lines = []
            while i < total:
                et, el, pi = elements[i]
                if et == 'p' and get_para_style(el) == '代码清单':
                    code_lines.append(escape_code_line(get_para_text(el)))
                    i += 1
                else:
                    break
            o.append(f'<section style="{S_CODE_WRAP}">')
            for cl in code_lines:
                o.append(f'<p style="{S_CODE_LINE}">{cl}</p>')
            o.append('</section>')
            continue

        # Bullet lists
        if text and (text.startswith('●') or text.startswith('•')):
            items = []
            while i < total:
                et, el, pi = elements[i]
                if et != 'p' or pi >= end_idx:
                    break
                t = get_para_text(el).strip()
                if t and (t.startswith('●') or t.startswith('•')):
                    items.append(escape(t.lstrip('●• ').strip()))
                    i += 1
                else:
                    break
            o.append('<ul style="margin:14px 0;padding-left:2em;">')
            for item in items:
                o.append(f'<li style="{S_LI}">{item}</li>')
            o.append('</ul>')
            continue

        # Image
        img_src = extract_image_base64(elem, doc_part)
        if img_src:
            caption = text
            o.append(f'<section style="{S_IW}">')
            o.append(f'<img style="{S_IMG}" src="{img_src}" />')
            if caption:
                o.append(f'<p style="{S_IC}">{escape(caption)}</p>')
            o.append('</section>')
            i += 1
            continue

        # Skip empty
        if not text:
            i += 1
            continue

        rich = build_rich_text_from_element(elem)
        style_lower = style.lower()

        if style_lower == 'heading 1':
            o.append(f'<section style="{S_H1}">{rich}</section>')
        elif style_lower == 'heading 2':
            o.append(f'<section style="{S_SEP}">- - -</section>')
            o.append(f'<section style="{S_H2}">{rich}</section>')
        elif style_lower == 'heading 3':
            o.append(f'<section style="{S_H3}">{rich}</section>')
        elif style_lower == 'heading 4':
            o.append(f'<section style="{S_H4}">{rich}</section>')
        elif style == '注意内容':
            o.append(f'<section style="{S_NOTE}">{rich}</section>')
        elif style == '插图插表':
            o.append(f'<p style="{S_IC}">{rich}</p>')
        else:
            o.append(f'<p style="{S_P}">{rich}</p>')

        i += 1

    # Footer
    o.append(f'<section style="{S_SEP}">- - -</section>')
    o.append(f'<section style="{S_FOOTER}">')
    o.append(f'本文摘自《AI大模型实战：从零基础到项目落地》<br>')
    o.append(f'觉得有帮助？点个 <strong style="{S_STRONG}">「在看」</strong> 支持一下吧')
    o.append('</section>')
    o.append('</section>')

    body = '\n'.join(o)
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{escape(title)}</title>
{WECHAT_CSS}
</head>
<body>
{body}
</body>
</html>"""


def main():
    sys.stdout.reconfigure(encoding='utf-8')

    doc_path = Path('Part1_基础篇.docx')
    if not doc_path.exists():
        print(f"Error: {doc_path} not found")
        return

    doc = Document(str(doc_path))
    output_dir = Path('wechat_articles')
    output_dir.mkdir(exist_ok=True)

    print(f"文档共 {len(doc.paragraphs)} 个段落, {len(doc.tables)} 个表格\n")

    for idx, article in enumerate(ARTICLES, 1):
        print(f"[{idx}/{len(ARTICLES)}] {article['title']}")
        html = convert_chapter(
            doc, article['start'], article['end'],
            article['title'], article['subtitle'],
        )
        out_path = output_dir / article['file']
        out_path.write_text(html, encoding='utf-8')
        size_kb = out_path.stat().st_size / 1024
        print(f"     -> {out_path.name} ({size_kb:.0f} KB)\n")

    print(f"完成！共 {len(ARTICLES)} 篇，保存在 {output_dir}/")
    print("使用方法：浏览器打开 HTML → Ctrl+A 全选 → Ctrl+C 复制 → 粘贴到微信公众号编辑器")


if __name__ == '__main__':
    main()
