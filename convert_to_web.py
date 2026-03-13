"""
将全部 5 个 Part.docx 转换为适合在线阅读的网页版 HTML
输出到 web_articles/ 目录，与微信版 wechat_articles/ 并行
"""
import sys
import re
import base64
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image

# ============================================================
# 全部 41 篇文章定义
# ============================================================
ARTICLES = [
    # Part1 基础篇
    {"docx": "Part1_基础篇.docx",  "file": "01.html", "num": "01", "part": "基础篇",
     "title": "AI大模型到底是什么？一篇文章带你从零搞懂",
     "subtitle": "从猜水果到GPT-4，人工智能的三次浪潮与大模型的前世今生",
     "start": 2, "end": 103},
    {"docx": "Part1_基础篇.docx",  "file": "02.html", "num": "02", "part": "基础篇",
     "title": "5分钟，让你亲手调用一次大模型API",
     "subtitle": "零代码经验也能跑通的AI实战第一课",
     "start": 103, "end": 194},
    {"docx": "Part1_基础篇.docx",  "file": "03.html", "num": "03", "part": "基础篇",
     "title": "AI开发第一步：Python环境搭建与语法速成",
     "subtitle": "Conda环境、变量类型、字符串操作、条件循环，一文搞定",
     "start": 194, "end": 423},
    {"docx": "Part1_基础篇.docx",  "file": "04.html", "num": "04", "part": "基础篇",
     "title": "Python进阶：函数、类、NumPy、Pandas一网打尽",
     "subtitle": "从模块化编程到数据处理，AI开发者的必备技能",
     "start": 423, "end": 865},
    {"docx": "Part1_基础篇.docx",  "file": "05.html", "num": "05", "part": "基础篇",
     "title": "深度学习没那么难：从神经网络到PyTorch实战",
     "subtitle": "感知机、损失函数、反向传播、CNN/RNN、GPU配置与手写数字识别",
     "start": 865, "end": 1180},
    {"docx": "Part1_基础篇.docx",  "file": "06.html", "num": "06", "part": "基础篇",
     "title": "让机器读懂人话：NLP基础与词嵌入技术详解",
     "subtitle": "分词、TF-IDF、Word2Vec、语言模型演进，一文讲透",
     "start": 1180, "end": 1326},
    {"docx": "Part1_基础篇.docx",  "file": "07.html", "num": "07", "part": "基础篇",
     "title": "手把手教你用HuggingFace做中文情感分析",
     "subtitle": "从Transformers库入门到完整实战项目，附完整代码",
     "start": 1326, "end": 99999},
    # Part2 核心篇
    {"docx": "Part2_核心篇.docx",  "file": "08.html", "num": "08", "part": "核心篇",
     "title": "Transformer为什么能统治AI？注意力机制全解析",
     "subtitle": "Attention、Self-Attention、位置编码、Encoder-Decoder，一文吃透核心原理",
     "start": 2, "end": 152},
    {"docx": "Part2_核心篇.docx",  "file": "09.html", "num": "09", "part": "核心篇",
     "title": "手撕Transformer：从零用PyTorch实现完整代码",
     "subtitle": "500行代码带你搞懂注意力机制的每一个细节",
     "start": 152, "end": 526},
    {"docx": "Part2_核心篇.docx",  "file": "10.html", "num": "10", "part": "核心篇",
     "title": "ChatGPT是怎么炼成的？大模型训练全流程揭秘",
     "subtitle": "预训练、分词器、Scaling Laws、SFT、RLHF、DPO，一文看懂",
     "start": 526, "end": 815},
    {"docx": "Part2_核心篇.docx",  "file": "11.html", "num": "11", "part": "核心篇",
     "title": "和AI说话也有技巧：Prompt Engineering入门指南",
     "subtitle": "角色设定、思维链、Few-Shot、结构化输出，让大模型听话干活",
     "start": 815, "end": 991},
    {"docx": "Part2_核心篇.docx",  "file": "12.html", "num": "12", "part": "核心篇",
     "title": "收藏级！高质量Prompt模板库，拿来就能用",
     "subtitle": "代码生成、数据分析、写作润色、翻译校对，覆盖10大场景",
     "start": 991, "end": 1230},
    {"docx": "Part2_核心篇.docx",  "file": "13.html", "num": "13", "part": "核心篇",
     "title": "大模型API怎么调？从认证到流式输出全攻略",
     "subtitle": "OpenAI/Claude/通义千问/DeepSeek四大平台API对比与实战",
     "start": 1230, "end": 1437},
    {"docx": "Part2_核心篇.docx",  "file": "14.html", "num": "14", "part": "核心篇",
     "title": "让AI调用你的代码：Function Calling工具调用详解",
     "subtitle": "从原理到实战，教大模型学会查天气、算数学、搜数据库",
     "start": 1437, "end": 1554},
    {"docx": "Part2_核心篇.docx",  "file": "15.html", "num": "15", "part": "核心篇",
     "title": "实战：开发一个多模型切换的智能助手",
     "subtitle": "支持GPT/Claude/DeepSeek自由切换，附完整可运行代码",
     "start": 1554, "end": 99999},
    # Part3 进阶篇
    {"docx": "Part3_进阶篇.docx",  "file": "16.html", "num": "16", "part": "进阶篇",
     "title": "大模型微调到底怎么做？LoRA/QLoRA原理全解析",
     "subtitle": "为什么需要微调、PEFT参数高效微调、LoRA数学直觉、显存对比，一文搞懂",
     "start": 2, "end": 150},
    {"docx": "Part3_进阶篇.docx",  "file": "17.html", "num": "17", "part": "进阶篇",
     "title": "实战！用QLoRA微调Qwen模型打造领域客服",
     "subtitle": "数据准备、超参数设置、训练监控、BLEU/ROUGE评估，附完整可运行代码",
     "start": 150, "end": 456},
    {"docx": "Part3_进阶篇.docx",  "file": "18.html", "num": "18", "part": "进阶篇",
     "title": "让大模型拥有外部知识：RAG检索增强生成全解析",
     "subtitle": "文档解析、文本切分策略、Embedding模型选型，从原理到代码一步步拆解",
     "start": 456, "end": 708},
    {"docx": "Part3_进阶篇.docx",  "file": "19.html", "num": "19", "part": "进阶篇",
     "title": "Milvus/Chroma/FAISS怎么选？向量库对比+企业知识库实战",
     "subtitle": "混合检索、Reranker重排序、RAG评估体系，附Streamlit界面完整代码",
     "start": 708, "end": 1173},
    {"docx": "Part3_进阶篇.docx",  "file": "20.html", "num": "20", "part": "进阶篇",
     "title": "LangChain从入门到精通：Model I/O、Chain、Memory全解析",
     "subtitle": "LCEL表达式语言、Prompt Templates、Output Parsers、对话记忆管理",
     "start": 1173, "end": 1457},
    {"docx": "Part3_进阶篇.docx",  "file": "21.html", "num": "21", "part": "进阶篇",
     "title": "LangSmith调试+LangGraph状态机，搭建复杂AI工作流",
     "subtitle": "可观测性追踪、条件分支工作流、多步骤文档分析助手，附完整代码",
     "start": 1457, "end": 1843},
    {"docx": "Part3_进阶篇.docx",  "file": "22.html", "num": "22", "part": "进阶篇",
     "title": "AI Agent到底是什么？ReAct模式与工具调用深度解析",
     "subtitle": "感知-规划-执行-反思架构、ReAct Prompt模板、搜索/代码/数据库工具集成",
     "start": 1843, "end": 2017},
    {"docx": "Part3_进阶篇.docx",  "file": "23.html", "num": "23", "part": "进阶篇",
     "title": "多Agent协作+MCP协议：构建具备联网搜索与代码执行能力的Agent",
     "subtitle": "LangGraph多Agent系统、长期记忆实现、MCP Server开发，附完整实战代码",
     "start": 2017, "end": 99999},
    # Part4 应用篇
    {"docx": "Part4_应用篇.docx",  "file": "24.html", "num": "24", "part": "应用篇",
     "title": "大模型「五感觉醒」：多模态技术全解析",
     "subtitle": "VLM图像理解、Whisper语音识别、Stable Diffusion文生图、视频生成，一文看懂",
     "start": 2, "end": 205},
    {"docx": "Part4_应用篇.docx",  "file": "25.html", "num": "25", "part": "应用篇",
     "title": "实战！用Qwen2.5-VL+Whisper构建多模态内容分析平台",
     "subtitle": "图文音频一体化处理，附Gradio界面完整代码",
     "start": 205, "end": 400},
    {"docx": "Part4_应用篇.docx",  "file": "26.html", "num": "26", "part": "应用篇",
     "title": "大模型应用怎么设计？架构、FastAPI服务、对话系统全攻略",
     "subtitle": "三层架构设计、流式接口、Gradio/Streamlit界面、多轮对话意图识别",
     "start": 400, "end": 800},
    {"docx": "Part4_应用篇.docx",  "file": "27.html", "num": "27", "part": "应用篇",
     "title": "实战：合同审核+代码助手+端到端智能写作助手",
     "subtitle": "文档处理流水线、代码生成Prompt设计、FastAPI+前端完整工程代码",
     "start": 800, "end": 1176},
    {"docx": "Part4_应用篇.docx",  "file": "28.html", "num": "28", "part": "应用篇",
     "title": "大模型推理太慢太贵？量化+推理框架帮你省钱提速",
     "subtitle": "INT4/INT8/GPTQ/AWQ量化对比、vLLM/TGI/Ollama框架选型、KV Cache优化、多GPU并行",
     "start": 1176, "end": 1351},
    {"docx": "Part4_应用篇.docx",  "file": "29.html", "num": "29", "part": "应用篇",
     "title": "手机上跑大模型？边缘部署+成本估算+vLLM生产级部署实战",
     "subtitle": "端侧模型推荐、GPU选型矩阵、云服务对比、Nginx负载均衡，附性能测试脚本",
     "start": 1351, "end": 1578},
    {"docx": "Part4_应用篇.docx",  "file": "30.html", "num": "30", "part": "应用篇",
     "title": "大模型怎么评分？MMLU、C-Eval、HumanEval评测基准全解析",
     "subtitle": "通用与领域评测基准、LLM-as-Judge、自动化评测流水线搭建",
     "start": 1578, "end": 1864},
    {"docx": "Part4_应用篇.docx",  "file": "31.html", "num": "31", "part": "应用篇",
     "title": "大模型会说谎、有偏见？幻觉、偏见、内容安全与红队测试详解",
     "subtitle": "安全风险分类、幻觉检测、内容过滤多层防护、自动化红队测试框架",
     "start": 1864, "end": 2253},
    {"docx": "Part4_应用篇.docx",  "file": "32.html", "num": "32", "part": "应用篇",
     "title": "实战：搭建一套完整的大模型评测与安全审计系统",
     "subtitle": "统一配置管理、评测+安全双引擎、一键运行完整项目代码",
     "start": 2253, "end": 99999},
    # Part5 实战篇
    {"docx": "Part5_实战篇.docx",  "file": "33.html", "num": "33", "part": "实战篇",
     "title": "从0搭一套智能客服系统：架构设计与RAG知识库管道",
     "subtitle": "C4架构全景图、技术选型矩阵、文档摄入切分策略，附核心代码实现",
     "start": 2, "end": 449},
    {"docx": "Part5_实战篇.docx",  "file": "34.html", "num": "34", "part": "实战篇",
     "title": "智能客服的大脑：双层意图识别+对话状态机+工单人机协作",
     "subtitle": "意图分类体系、对话状态转换图、工单路由决策引擎，完整代码实现",
     "start": 449, "end": 1037},
    {"docx": "Part5_实战篇.docx",  "file": "35.html", "num": "35", "part": "实战篇",
     "title": "智能客服上线全攻略：A/B测试+Docker部署+运维监控",
     "subtitle": "性能优化决策树、A/B测试框架、Docker Compose配置、FastAPI主应用",
     "start": 1037, "end": 1500},
    {"docx": "Part5_实战篇.docx",  "file": "36.html", "num": "36", "part": "实战篇",
     "title": "打造垂直领域大模型：数据工程五层漏斗+继续预训练实战",
     "subtitle": "通用vs垂直模型对比、数据清洗标注全流程、CPT训练代码完整实现",
     "start": 1500, "end": 1945},
    {"docx": "Part5_实战篇.docx",  "file": "37.html", "num": "37", "part": "实战篇",
     "title": "医疗/法律/金融大模型实践：LoRA微调+评估+数据飞轮闭环",
     "subtitle": "SFT数据格式、三大领域对比、难例挖掘策略、版本迭代管理，附避坑指南",
     "start": 1945, "end": 2413},
    {"docx": "Part5_实战篇.docx",  "file": "38.html", "num": "38", "part": "实战篇",
     "title": "单Agent不够用了？多Agent协作架构模式与DAG工作流编排",
     "subtitle": "Agent能力演进路线图、协作拓扑对比、有向无环图引擎核心实现",
     "start": 2413, "end": 2709},
    {"docx": "Part5_实战篇.docx",  "file": "39.html", "num": "39", "part": "实战篇",
     "title": "规划者+执行者+审核者：多Agent角色设计与工具链集成详解",
     "subtitle": "角色关系架构图、Agent通信协议、API/数据库/文件系统工具注册框架",
     "start": 2709, "end": 3288},
    {"docx": "Part5_实战篇.docx",  "file": "40.html", "num": "40", "part": "实战篇",
     "title": "实战！多Agent自动生成数据分析报告：异常处理+人工介入+完整编排",
     "subtitle": "异常处理中间件、端到端系统流程图、完整编排器实现，附架构演进路线",
     "start": 3288, "end": 3757},
    {"docx": "Part5_实战篇.docx",  "file": "41.html", "num": "41", "part": "实战篇",
     "title": "大模型未来在哪里？推理模型、MoE、具身智能+从入门到专家学习路线",
     "subtitle": "前沿趋势全景、小模型逆袭、AI创业商业模式、必读论文+开源项目推荐清单",
     "start": 3757, "end": 99999},
]

# ============================================================
# 常量
# ============================================================
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
IMG_MAX_WIDTH = 1200
IMG_JPEG_QUALITY = 85

PART_COLORS = {
    "基础篇": "#2563eb",
    "核心篇": "#7c3aed",
    "进阶篇": "#059669",
    "应用篇": "#d97706",
    "实战篇": "#dc2626",
}

# ============================================================
# CSS
# ============================================================
WEB_CSS = """\
:root {
  --accent: #2563eb;
  --accent-light: #eff6ff;
  --text: #1e293b;
  --text-muted: #64748b;
  --border: #e2e8f0;
  --code-bg: #1e293b;
  --code-text: #e2e8f0;
  --note-bg: #f0fdf4;
  --note-border: #22c55e;
  --radius: 8px;
  --max-width: 820px;
}
* { margin: 0; padding: 0; box-sizing: border-box; }
html { scroll-behavior: smooth; }
body {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC",
               "Hiragino Sans GB", "Microsoft YaHei", sans-serif;
  font-size: 17px; line-height: 1.85; color: var(--text);
  background: #f8fafc;
}

/* ---- 顶部导航 ---- */
.site-header {
  position: sticky; top: 0; z-index: 100;
  background: #fff; border-bottom: 1px solid var(--border);
  padding: 0 24px;
}
.site-header-inner {
  max-width: var(--max-width); margin: 0 auto;
  height: 52px; display: flex; align-items: center; gap: 12px;
}
.site-title { font-weight: 700; font-size: 15px; color: var(--text); text-decoration: none; }
.site-title:hover { color: var(--accent); }
.header-sep { color: var(--border); }
.part-badge {
  font-size: 12px; font-weight: 600; padding: 2px 10px;
  border-radius: 20px; color: #fff;
}

/* ---- 主体 ---- */
.page-wrapper { max-width: var(--max-width); margin: 0 auto; padding: 48px 24px 80px; }

/* ---- 文章头 ---- */
.article-header { margin-bottom: 40px; }
.article-num { font-size: 13px; color: var(--text-muted); margin-bottom: 10px; }
.article-title { font-size: 28px; font-weight: 800; line-height: 1.35; color: var(--text); margin-bottom: 12px; }
.article-subtitle { font-size: 15px; color: var(--text-muted); line-height: 1.6; padding-bottom: 28px; border-bottom: 2px solid var(--border); }

/* ---- 正文 ---- */
.article-body p { margin-bottom: 20px; text-align: justify; }

h2 { font-size: 22px; font-weight: 700; color: var(--text); margin: 44px 0 16px;
     padding-left: 14px; border-left: 4px solid var(--accent); }
h3 { font-size: 19px; font-weight: 700; color: var(--text); margin: 32px 0 12px; }
h4 { font-size: 17px; font-weight: 700; color: #334155; margin: 24px 0 10px; }
h5 { font-size: 16px; font-weight: 600; color: #475569; margin: 20px 0 8px; }

/* ---- 注意块 ---- */
.note-block {
  background: var(--note-bg); border-left: 4px solid var(--note-border);
  padding: 16px 20px; margin: 24px 0; border-radius: 0 var(--radius) var(--radius) 0;
  font-size: 15px; color: #166534; line-height: 1.7;
}

/* ---- 代码块 ---- */
.code-block {
  background: var(--code-bg); border-radius: var(--radius);
  padding: 20px 24px; margin: 24px 0; overflow-x: auto;
  -webkit-overflow-scrolling: touch;
}
.code-block pre {
  font-family: "JetBrains Mono", "Fira Code", Menlo, Consolas, monospace;
  font-size: 13.5px; line-height: 1.65; color: var(--code-text);
  white-space: pre; margin: 0;
}

/* ---- 表格 ---- */
.table-wrap { overflow-x: auto; margin: 24px 0; border-radius: var(--radius); border: 1px solid var(--border); }
table { width: 100%; border-collapse: collapse; font-size: 14px; }
th { background: var(--accent); color: #fff; font-weight: 600; padding: 11px 14px; text-align: left; }
td { padding: 10px 14px; border-top: 1px solid var(--border); color: var(--text); line-height: 1.6; }
tr:nth-child(even) td { background: #f8fafc; }

/* ---- 图片 ---- */
.img-wrap { text-align: center; margin: 28px 0; }
.img-wrap img { max-width: 100%; height: auto; border-radius: var(--radius); box-shadow: 0 4px 20px rgba(0,0,0,0.08); }
.img-caption { font-size: 13px; color: var(--text-muted); margin-top: 10px; }

/* ---- 列表 ---- */
ul, ol { margin: 16px 0 20px 1.6em; }
li { margin-bottom: 8px; line-height: 1.75; }

/* ---- 图表题注 ---- */
.caption { font-size: 13px; color: var(--text-muted); text-align: center; margin: -16px 0 20px; }

/* ---- strong / em ---- */
strong { color: #1d4ed8; font-weight: 700; }
em { font-style: normal; color: #b45309; border-bottom: 1px dashed #b45309; }

/* ---- 上下篇导航 ---- */
.article-nav {
  display: flex; justify-content: space-between; gap: 16px;
  max-width: var(--max-width); margin: 0 auto; padding: 0 24px 60px;
}
.nav-btn {
  flex: 1; display: flex; flex-direction: column; padding: 16px 20px;
  background: #fff; border: 1px solid var(--border); border-radius: var(--radius);
  text-decoration: none; color: var(--text); transition: border-color .15s, box-shadow .15s;
  min-width: 0;
}
.nav-btn:hover { border-color: var(--accent); box-shadow: 0 2px 8px rgba(37,99,235,.1); }
.nav-btn.disabled { opacity: .4; pointer-events: none; }
.nav-label { font-size: 12px; color: var(--text-muted); margin-bottom: 4px; }
.nav-title { font-size: 14px; font-weight: 600; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
.nav-btn.next { text-align: right; }

/* ---- 底部 ---- */
.site-footer {
  background: #fff; border-top: 1px solid var(--border);
  padding: 40px 24px; text-align: center; color: var(--text-muted); font-size: 14px;
}
.footer-inner { max-width: var(--max-width); margin: 0 auto; }
.footer-qr img { width: 130px; height: 130px; border-radius: var(--radius); margin-bottom: 12px; }
.footer-title { font-size: 16px; font-weight: 700; color: var(--text); margin-bottom: 6px; }
.footer-desc { line-height: 1.7; }

/* ---- 响应式 ---- */
@media (max-width: 640px) {
  .page-wrapper { padding: 28px 16px 60px; }
  .article-title { font-size: 22px; }
  h2 { font-size: 19px; }
  h3 { font-size: 17px; }
  .article-nav { flex-direction: column; padding: 0 16px 40px; }
  .code-block { padding: 14px 16px; }
  .code-block pre { font-size: 12.5px; }
}
"""


# ============================================================
# 工具函数
# ============================================================
def escape(text: str) -> str:
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def compress_image_to_base64(blob: bytes, content_type: str) -> str:
    try:
        img = Image.open(BytesIO(blob))
    except Exception:
        return f"data:{content_type};base64,{base64.b64encode(blob).decode()}"
    if img.width > IMG_MAX_WIDTH:
        ratio = IMG_MAX_WIDTH / img.width
        img = img.resize((IMG_MAX_WIDTH, int(img.height * ratio)), Image.LANCZOS)
    has_alpha = img.mode == 'RGBA' and img.getchannel('A').getextrema()[0] < 250
    buf = BytesIO()
    if has_alpha:
        img.save(buf, format='PNG', optimize=True)
        mime = 'image/png'
    else:
        img.convert('RGB').save(buf, format='JPEG', quality=IMG_JPEG_QUALITY, optimize=True)
        mime = 'image/jpeg'
    return f"data:{mime};base64,{base64.b64encode(buf.getvalue()).decode()}"


def extract_image_base64(elem, doc_part):
    blips = elem.findall(f'.//{{{A_NS}}}blip')
    if not blips:
        return None
    embed_id = blips[0].get(f'{{{R_NS}}}embed')
    if not embed_id:
        return None
    try:
        ip = doc_part.related_parts[embed_id]
        return compress_image_to_base64(ip.blob, ip.content_type)
    except (KeyError, AttributeError):
        return None


def get_para_text(elem) -> str:
    return ''.join(t.text or '' for t in elem.iter(qn('w:t')))


def get_para_style(elem) -> str:
    pPr = elem.find(qn('w:pPr'))
    if pPr is not None:
        ps = pPr.find(qn('w:pStyle'))
        if ps is not None:
            return ps.get(qn('w:val'), 'Normal')
    return 'Normal'


def build_rich_text(elem) -> str:
    parts = []
    for run in elem.findall(qn('w:r')):
        t = ''.join((te.text or '') for te in run.findall(qn('w:t')))
        if not t:
            continue
        t = escape(t)
        rPr = run.find(qn('w:rPr'))
        bold = italic = False
        if rPr is not None:
            b = rPr.find(qn('w:b'))
            if b is not None and b.get(qn('w:val'), 'true') != 'false':
                bold = True
            i = rPr.find(qn('w:i'))
            if i is not None and i.get(qn('w:val'), 'true') != 'false':
                italic = True
        if bold and italic:
            t = f'<strong><em>{t}</em></strong>'
        elif bold:
            t = f'<strong>{t}</strong>'
        elif italic:
            t = f'<em>{t}</em>'
        parts.append(t)
    result = ''.join(parts)
    return result if result else escape(get_para_text(elem))


def build_body_elements(doc):
    body = doc.element.body
    items, para_idx = [], 0
    for child in body:
        tag = etree.QName(child).localname
        if tag == 'p':
            items.append(('p', child, para_idx))
            para_idx += 1
        elif tag == 'tbl':
            items.append(('tbl', child, para_idx))
    return items


def table_to_html(tbl_elem, doc) -> str:
    table = Table(tbl_elem, doc)
    if not table.rows:
        return ''
    rows_html = []
    for r_idx, row in enumerate(table.rows):
        cells_html = []
        for cell in row.cells:
            texts = [escape(p.text.strip()) for p in cell.paragraphs if p.text.strip()]
            content = '<br>'.join(texts) or '&nbsp;'
            tag = 'th' if r_idx == 0 else 'td'
            cells_html.append(f'<{tag}>{content}</{tag}>')
        rows_html.append('<tr>' + ''.join(cells_html) + '</tr>')
    return f'<div class="table-wrap"><table>{"".join(rows_html)}</table></div>'


# ============================================================
# 核心转换函数
# ============================================================
def convert_chapter_web(doc, start_idx: int, end_idx: int) -> str:
    doc_part = doc.part
    elements = build_body_elements(doc)
    end_idx = min(end_idx, len(doc.paragraphs))
    o = []
    i = 0
    total = len(elements)

    while i < total:
        etype, elem, pidx = elements[i]

        if etype == 'tbl':
            if start_idx <= pidx < end_idx:
                o.append(table_to_html(elem, doc))
            i += 1
            continue

        if pidx < start_idx:
            i += 1
            continue
        if pidx >= end_idx:
            break

        style = get_para_style(elem)
        text = get_para_text(elem).strip()

        # 代码块：合并连续行，用 <pre> 保留缩进
        if style == '代码清单':
            code_lines = []
            while i < total:
                et, el, pi = elements[i]
                if et == 'p' and pi < end_idx and get_para_style(el) == '代码清单':
                    code_lines.append(escape(get_para_text(el)))
                    i += 1
                else:
                    break
            code = '\n'.join(code_lines)
            o.append(f'<div class="code-block"><pre>{code}</pre></div>')
            continue

        # 图片
        img_src = extract_image_base64(elem, doc_part)
        if img_src:
            caption = escape(text) if text else ''
            o.append('<div class="img-wrap">')
            o.append(f'<img src="{img_src}" loading="lazy" />')
            if caption:
                o.append(f'<p class="img-caption">{caption}</p>')
            o.append('</div>')
            i += 1
            continue

        # 空段落跳过
        if not text:
            i += 1
            continue

        rich = build_rich_text(elem)
        sl = style.lower()

        if sl == 'heading 1':
            o.append(f'<h2>{rich}</h2>')
        elif sl == 'heading 2':
            o.append(f'<h3>{rich}</h3>')
        elif sl == 'heading 3':
            o.append(f'<h4>{rich}</h4>')
        elif sl == 'heading 4':
            o.append(f'<h5>{rich}</h5>')
        elif style == '注意内容':
            o.append(f'<div class="note-block">{rich}</div>')
        elif style == '插图插表':
            o.append(f'<p class="caption">{rich}</p>')
        elif text.startswith('●') or text.startswith('•'):
            # 收集连续列表项
            items = []
            while i < total:
                et, el, pi = elements[i]
                if et != 'p' or pi >= end_idx:
                    break
                t = get_para_text(el).strip()
                if t and (t.startswith('●') or t.startswith('•')):
                    items.append(build_rich_text(el).lstrip('●•&nbsp; ').strip())
                    i += 1
                else:
                    break
            o.append('<ul>' + ''.join(f'<li>{it}</li>' for it in items) + '</ul>')
            continue
        else:
            o.append(f'<p>{rich}</p>')

        i += 1

    return '\n'.join(o)


# ============================================================
# 页面组装
# ============================================================
def make_page(article: dict, body_html: str, prev_a, next_a, accent: str) -> str:
    num = article['num']
    title = escape(article['title'])
    subtitle = escape(article['subtitle'])
    part = escape(article['part'])

    # 上下篇导航
    prev_btn = (
        f'<a class="nav-btn prev" href="{prev_a["file"]}">'
        f'<span class="nav-label">← 上一篇</span>'
        f'<span class="nav-title">{escape(prev_a["title"])}</span></a>'
        if prev_a else
        '<a class="nav-btn prev disabled"><span class="nav-label">← 上一篇</span><span class="nav-title">已是第一篇</span></a>'
    )
    next_btn = (
        f'<a class="nav-btn next" href="{next_a["file"]}">'
        f'<span class="nav-label">下一篇 →</span>'
        f'<span class="nav-title">{escape(next_a["title"])}</span></a>'
        if next_a else
        '<a class="nav-btn next disabled"><span class="nav-label">下一篇 →</span><span class="nav-title">已是最后一篇</span></a>'
    )

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} · AI大模型实战</title>
<style>
{WEB_CSS}
:root {{ --accent: {accent}; }}
</style>
</head>
<body>

<header class="site-header">
  <div class="site-header-inner">
    <a class="site-title" href="index.html">AI大模型实战</a>
    <span class="header-sep">/</span>
    <span class="part-badge" style="background:{accent}">{part}</span>
    <span style="font-size:13px;color:#94a3b8;margin-left:auto">第 {num} 篇 / 共 41 篇</span>
  </div>
</header>

<main class="page-wrapper">
  <div class="article-header">
    <div class="article-num">第 {num} 篇 · {part}</div>
    <h1 class="article-title">{title}</h1>
    <p class="article-subtitle">{subtitle}</p>
  </div>
  <div class="article-body">
{body_html}
  </div>
</main>

<nav class="article-nav">
  {prev_btn}
  {next_btn}
</nav>

<footer class="site-footer">
  <div class="footer-inner">
    <div class="footer-qr">
      <img src="../assets/wechat_qrcode.jpg" alt="AI数学与未来" />
    </div>
    <div class="footer-title">关注公众号「AI 数学与未来」</div>
    <div class="footer-desc">持续更新 AI 大模型技术科普，从数学原理到工程落地<br>扫码关注，获取系列文章持续更新</div>
  </div>
</footer>

</body>
</html>"""


# ============================================================
# 首页
# ============================================================
def make_index() -> str:
    parts_order = ["基础篇", "核心篇", "进阶篇", "应用篇", "实战篇"]
    part_desc = {
        "基础篇": "大模型概念、Python基础、深度学习入门、NLP基础（共7篇）",
        "核心篇": "Transformer原理、大模型训练、Prompt工程、API调用（共8篇）",
        "进阶篇": "Fine-Tuning微调、RAG检索增强、LangChain、AI Agent（共8篇）",
        "应用篇": "多模态应用、应用开发、推理部署优化、安全评测（共9篇）",
        "实战篇": "智能客服、垂直大模型、多Agent工作流、技术趋势（共9篇）",
    }
    sections = []
    for part in parts_order:
        color = PART_COLORS[part]
        arts = [a for a in ARTICLES if a['part'] == part]
        cards = []
        for a in arts:
            cards.append(
                f'<a class="idx-card" href="web_articles/{a["file"]}">'
                f'<span class="idx-num" style="color:{color}">#{a["num"]}</span>'
                f'<span class="idx-title">{escape(a["title"])}</span>'
                f'</a>'
            )
        sections.append(
            f'<section class="idx-section">'
            f'<div class="idx-part-header">'
            f'<span class="idx-part-badge" style="background:{color}">{part}</span>'
            f'<span class="idx-part-desc">{part_desc[part]}</span>'
            f'</div>'
            f'<div class="idx-cards">{"".join(cards)}</div>'
            f'</section>'
        )

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>AI大模型实战：从零基础到项目落地</title>
<style>
{WEB_CSS}
.hero {{ background: linear-gradient(135deg,#1e40af,#7c3aed); color:#fff; padding:64px 24px; text-align:center; }}
.hero-inner {{ max-width:700px; margin:0 auto; }}
.hero h1 {{ font-size:32px; font-weight:800; margin-bottom:14px; line-height:1.3; }}
.hero p {{ font-size:17px; opacity:.85; line-height:1.7; }}
.hero-stats {{ display:flex; justify-content:center; gap:40px; margin-top:32px; }}
.stat {{ font-size:13px; opacity:.75; }}
.stat strong {{ display:block; font-size:26px; font-weight:800; opacity:1; margin-bottom:2px; }}
.idx-body {{ max-width:820px; margin:0 auto; padding:48px 24px 80px; }}
.idx-section {{ margin-bottom:44px; }}
.idx-part-header {{ display:flex; align-items:center; gap:14px; margin-bottom:16px; }}
.idx-part-badge {{ font-size:13px; font-weight:700; padding:4px 14px; border-radius:20px; color:#fff; }}
.idx-part-desc {{ font-size:14px; color:#64748b; }}
.idx-cards {{ display:grid; gap:10px; }}
.idx-card {{
  display:flex; align-items:flex-start; gap:12px; padding:14px 18px;
  background:#fff; border:1px solid #e2e8f0; border-radius:8px;
  text-decoration:none; color:#1e293b; transition:border-color .15s, box-shadow .15s;
}}
.idx-card:hover {{ border-color:var(--accent); box-shadow:0 2px 8px rgba(37,99,235,.1); }}
.idx-num {{ font-size:13px; font-weight:700; min-width:32px; padding-top:1px; }}
.idx-title {{ font-size:15px; line-height:1.5; }}
.footer-qr img {{ width:130px; border-radius:8px; margin-bottom:12px; }}
.footer-title {{ font-size:16px; font-weight:700; color:#1e293b; margin-bottom:6px; }}
</style>
</head>
<body>
<header class="site-header">
  <div class="site-header-inner">
    <a class="site-title" href="index.html">AI大模型实战</a>
  </div>
</header>

<div class="hero">
  <div class="hero-inner">
    <h1>AI大模型实战<br>从零基础到项目落地</h1>
    <p>系统讲解大模型技术栈，从 Python 基础、Transformer 原理，到 RAG、微调、Agent，再到生产级部署与安全评测</p>
    <div class="hero-stats">
      <div class="stat"><strong>41</strong>篇文章</div>
      <div class="stat"><strong>5</strong>大篇章</div>
      <div class="stat"><strong>20</strong>章节</div>
    </div>
  </div>
</div>

<div class="idx-body">
  {''.join(sections)}
</div>

<footer class="site-footer">
  <div class="footer-inner">
    <div class="footer-qr">
      <img src="assets/wechat_qrcode.jpg" alt="AI数学与未来" />
    </div>
    <div class="footer-title">关注公众号「AI 数学与未来」</div>
    <div class="footer-desc">持续更新 AI 大模型技术科普，从数学原理到工程落地<br>扫码关注，获取系列文章持续更新</div>
  </div>
</footer>
</body>
</html>"""


# ============================================================
# 主函数
# ============================================================
def main():
    sys.stdout.reconfigure(encoding='utf-8')
    out_dir = Path('web_articles')
    out_dir.mkdir(exist_ok=True)

    # 按 docx 分组，避免重复加载
    from itertools import groupby
    docs_cache = {}

    total = len(ARTICLES)
    for idx, article in enumerate(ARTICLES):
        docx_name = article['docx']
        if docx_name not in docs_cache:
            p = Path(docx_name)
            if not p.exists():
                print(f"  [跳过] {docx_name} 不存在")
                docs_cache[docx_name] = None
            else:
                print(f"\n加载 {docx_name} ...")
                docs_cache[docx_name] = Document(str(p))

        doc = docs_cache[docx_name]
        if doc is None:
            continue

        prev_a = ARTICLES[idx - 1] if idx > 0 else None
        next_a = ARTICLES[idx + 1] if idx < total - 1 else None
        accent = PART_COLORS.get(article['part'], '#2563eb')

        print(f"[{article['num']}/{total}] {article['title'][:50]}")
        body = convert_chapter_web(doc, article['start'], article['end'])
        html = make_page(article, body, prev_a, next_a, accent)

        out_path = out_dir / article['file']
        out_path.write_text(html, encoding='utf-8')
        size_kb = out_path.stat().st_size / 1024
        print(f"       -> {out_path.name}  ({size_kb:.0f} KB)")

    # 生成首页
    print("\n生成首页 index.html ...")
    Path('index.html').write_text(make_index(), encoding='utf-8')
    print("完成！")
    print(f"\n在线访问：https://shaoxueheng.github.io/AI-LLM-Practice/")


if __name__ == '__main__':
    main()
