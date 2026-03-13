"""
将 Part3_进阶篇.docx 转换为多篇微信公众号 HTML 文章
复用 convert_to_wechat.py 的转换引擎
"""
import sys
from pathlib import Path
from docx import Document

# 复用 Part1 脚本的所有转换函数
from convert_to_wechat import convert_chapter

# ============================================================
# Part3 文章拆分方案（8篇）
# ============================================================
ARTICLES = [
    # --- 第9章 大模型微调（拆为2篇） ---
    {
        "file": "16_大模型微调原理与LoRA技术.html",
        "title": "大模型微调到底怎么做？LoRA/QLoRA原理全解析",
        "subtitle": "为什么需要微调、PEFT参数高效微调、LoRA数学直觉、显存对比，一文搞懂",
        "start": 2, "end": 150,
    },
    {
        "file": "17_QLoRA微调Qwen实战.html",
        "title": "实战！用QLoRA微调Qwen模型打造领域客服",
        "subtitle": "数据准备、超参数设置、训练监控、BLEU/ROUGE评估，附完整可运行代码",
        "start": 150, "end": 456,
    },
    # --- 第10章 RAG 检索增强生成（拆为2篇） ---
    {
        "file": "18_RAG检索增强生成原理.html",
        "title": "让大模型拥有外部知识：RAG检索增强生成全解析",
        "subtitle": "文档解析、文本切分策略、Embedding模型选型，从原理到代码一步步拆解",
        "start": 456, "end": 708,
    },
    {
        "file": "19_向量数据库与RAG实战.html",
        "title": "Milvus/Chroma/FAISS怎么选？向量库对比+企业知识库实战",
        "subtitle": "混合检索、Reranker重排序、RAG评估体系，附Streamlit界面完整代码",
        "start": 708, "end": 1173,
    },
    # --- 第11章 LangChain（拆为2篇） ---
    {
        "file": "20_LangChain核心组件详解.html",
        "title": "LangChain从入门到精通：Model I/O、Chain、Memory全解析",
        "subtitle": "LCEL表达式语言、Prompt Templates、Output Parsers、对话记忆管理",
        "start": 1173, "end": 1457,
    },
    {
        "file": "21_LangGraph工作流与实战.html",
        "title": "LangSmith调试+LangGraph状态机，搭建复杂AI工作流",
        "subtitle": "可观测性追踪、条件分支工作流、多步骤文档分析助手，附完整代码",
        "start": 1457, "end": 1843,
    },
    # --- 第12章 AI Agent（拆为2篇） ---
    {
        "file": "22_AIAgent原理与ReAct模式.html",
        "title": "AI Agent到底是什么？ReAct模式与工具调用深度解析",
        "subtitle": "感知-规划-执行-反思架构、ReAct Prompt模板、搜索/代码/数据库工具集成",
        "start": 1843, "end": 2017,
    },
    {
        "file": "23_多Agent协作与MCP协议实战.html",
        "title": "多Agent协作+MCP协议：构建具备联网搜索与代码执行能力的Agent",
        "subtitle": "LangGraph多Agent系统、长期记忆实现、MCP Server开发，附完整实战代码",
        "start": 2017, "end": 99999,
    },
]


def main():
    sys.stdout.reconfigure(encoding='utf-8')

    doc_path = Path('Part3_进阶篇.docx')
    if not doc_path.exists():
        print(f"Error: {doc_path} not found")
        return

    doc = Document(str(doc_path))
    output_dir = Path('wechat_articles')
    output_dir.mkdir(exist_ok=True)

    print(f"文档共 {len(doc.paragraphs)} 个段落, {len(doc.tables)} 个表格\n")

    for idx, article in enumerate(ARTICLES, 1):
        print(f"[{idx}/{len(ARTICLES)}] {article['title']}")
        html = convert_chapter(
            doc, article['start'], article['end'],
            article['title'], article['subtitle'],
        )
        out_path = output_dir / article['file']
        out_path.write_text(html, encoding='utf-8')
        size_kb = out_path.stat().st_size / 1024
        print(f"     -> {out_path.name} ({size_kb:.0f} KB)\n")

    print(f"完成！共 {len(ARTICLES)} 篇，保存在 {output_dir}/")
    print("使用方法：浏览器打开 HTML → Ctrl+A 全选 → Ctrl+C 复制 → 粘贴到微信公众号编辑器")


if __name__ == '__main__':
    main()
