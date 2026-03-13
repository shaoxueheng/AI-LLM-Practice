"""
将 Part2_核心篇.docx 转换为多篇微信公众号 HTML 文章
复用 convert_to_wechat.py 的转换引擎
"""
import sys
from pathlib import Path
from docx import Document

# 复用 Part1 脚本的所有转换函数
from convert_to_wechat import convert_chapter

# ============================================================
# Part2 文章拆分方案（8篇）
# ============================================================
ARTICLES = [
    # --- 第5章 Transformer（拆为2篇） ---
    {
        "file": "08_Transformer原理深度解析.html",
        "title": "Transformer为什么能统治AI？注意力机制全解析",
        "subtitle": "Attention、Self-Attention、位置编码、Encoder-Decoder，一文吃透核心原理",
        "start": 2, "end": 152,
    },
    {
        "file": "09_从零实现Transformer.html",
        "title": "手撕Transformer：从零用PyTorch实现完整代码",
        "subtitle": "500行代码带你搞懂注意力机制的每一个细节",
        "start": 152, "end": 526,
    },
    # --- 第6章 大语言模型训练流程 ---
    {
        "file": "10_大模型训练全流程揭秘.html",
        "title": "ChatGPT是怎么炼成的？大模型训练全流程揭秘",
        "subtitle": "预训练、分词器、Scaling Laws、SFT、RLHF、DPO，一文看懂",
        "start": 526, "end": 815,
    },
    # --- 第7章 Prompt Engineering（拆为2篇） ---
    {
        "file": "11_提示词工程入门.html",
        "title": "和AI说话也有技巧：Prompt Engineering入门指南",
        "subtitle": "角色设定、思维链、Few-Shot、结构化输出，让大模型听话干活",
        "start": 815, "end": 991,
    },
    {
        "file": "12_Prompt模板库实战.html",
        "title": "收藏级！高质量Prompt模板库，拿来就能用",
        "subtitle": "代码生成、数据分析、写作润色、翻译校对，覆盖10大场景",
        "start": 991, "end": 1230,
    },
    # --- 第8章 大模型API调用（拆为3篇） ---
    {
        "file": "13_大模型API入门.html",
        "title": "大模型API怎么调？从认证到流式输出全攻略",
        "subtitle": "OpenAI/Claude/通义千问/DeepSeek四大平台API对比与实战",
        "start": 1230, "end": 1437,
    },
    {
        "file": "14_FunctionCalling工具调用.html",
        "title": "让AI调用你的代码：Function Calling工具调用详解",
        "subtitle": "从原理到实战，教大模型学会查天气、算数学、搜数据库",
        "start": 1437, "end": 1554,
    },
    {
        "file": "15_多模型智能助手实战.html",
        "title": "实战：开发一个多模型切换的智能助手",
        "subtitle": "支持GPT/Claude/DeepSeek自由切换，附完整可运行代码",
        "start": 1554, "end": 99999,
    },
]


def main():
    sys.stdout.reconfigure(encoding='utf-8')

    doc_path = Path('Part2_核心篇.docx')
    if not doc_path.exists():
        print(f"Error: {doc_path} not found")
        return

    doc = Document(str(doc_path))
    output_dir = Path('wechat_articles')
    output_dir.mkdir(exist_ok=True)

    print(f"文档共 {len(doc.paragraphs)} 个段落, {len(doc.tables)} 个表格\n")

    for idx, article in enumerate(ARTICLES, 1):
        print(f"[{idx}/{len(ARTICLES)}] {article['title']}")
        html = convert_chapter(
            doc, article['start'], article['end'],
            article['title'], article['subtitle'],
        )
        out_path = output_dir / article['file']
        out_path.write_text(html, encoding='utf-8')
        size_kb = out_path.stat().st_size / 1024
        print(f"     -> {out_path.name} ({size_kb:.0f} KB)\n")

    print(f"完成！共 {len(ARTICLES)} 篇，保存在 {output_dir}/")
    print("使用方法：浏览器打开 HTML → Ctrl+A 全选 → Ctrl+C 复制 → 粘贴到微信公众号编辑器")


if __name__ == '__main__':
    main()
