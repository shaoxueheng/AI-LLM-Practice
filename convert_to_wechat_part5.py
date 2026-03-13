"""
将 Part5_实战篇.docx 转换为多篇微信公众号 HTML 文章
复用 convert_to_wechat.py 的转换引擎
"""
import sys
from pathlib import Path
from docx import Document

# 复用 Part1 脚本的所有转换函数
from convert_to_wechat import convert_chapter

# ============================================================
# Part5 文章拆分方案（9篇）
# ============================================================
ARTICLES = [
    # --- 第17章 智能客服系统（拆为3篇） ---
    {
        "file": "33_智能客服系统架构与RAG管道.html",
        "title": "从0搭一套智能客服系统：架构设计与RAG知识库管道",
        "subtitle": "C4架构全景图、技术选型矩阵、文档摄入切分策略，附核心代码实现",
        "start": 2, "end": 449,
    },
    {
        "file": "34_智能客服意图识别与对话管理.html",
        "title": "智能客服的大脑：双层意图识别+对话状态机+工单人机协作",
        "subtitle": "意图分类体系、对话状态转换图、工单路由决策引擎，完整代码实现",
        "start": 449, "end": 1037,
    },
    {
        "file": "35_智能客服性能调优与生产部署.html",
        "title": "智能客服上线全攻略：A/B测试+Docker部署+运维监控",
        "subtitle": "性能优化决策树、A/B测试框架、Docker Compose配置、FastAPI主应用",
        "start": 1037, "end": 1500,
    },
    # --- 第18章 行业垂直大模型（拆为2篇） ---
    {
        "file": "36_垂直大模型数据工程与继续预训练.html",
        "title": "打造垂直领域大模型：数据工程五层漏斗+继续预训练实战",
        "subtitle": "通用vs垂直模型对比、数据清洗标注全流程、CPT训练代码完整实现",
        "start": 1500, "end": 1945,
    },
    {
        "file": "37_垂直大模型微调案例与数据飞轮.html",
        "title": "医疗/法律/金融大模型实践：LoRA微调+评估+数据飞轮闭环",
        "subtitle": "SFT数据格式、三大领域对比、难例挖掘策略、版本迭代管理，附避坑指南",
        "start": 1945, "end": 2413,
    },
    # --- 第19章 多Agent自动化工作流（拆为3篇） ---
    {
        "file": "38_多Agent工作流场景与DAG编排.html",
        "title": "单Agent不够用了？多Agent协作架构模式与DAG工作流编排",
        "subtitle": "Agent能力演进路线图、协作拓扑对比、有向无环图引擎核心实现",
        "start": 2413, "end": 2709,
    },
    {
        "file": "39_多Agent角色设计与工具链集成.html",
        "title": "规划者+执行者+审核者：多Agent角色设计与工具链集成详解",
        "subtitle": "角色关系架构图、Agent通信协议、API/数据库/文件系统工具注册框架",
        "start": 2709, "end": 3288,
    },
    {
        "file": "40_多Agent异常处理与数据分析报告实战.html",
        "title": "实战！多Agent自动生成数据分析报告：异常处理+人工介入+完整编排",
        "subtitle": "异常处理中间件、端到端系统流程图、完整编排器实现，附架构演进路线",
        "start": 3288, "end": 3757,
    },
    # --- 第20章 技术趋势与学习路线（1篇） ---
    {
        "file": "41_大模型技术趋势与学习路线.html",
        "title": "大模型未来在哪里？推理模型、MoE、具身智能+从入门到专家学习路线",
        "subtitle": "前沿趋势全景、小模型逆袭、AI创业商业模式、必读论文+开源项目推荐清单",
        "start": 3757, "end": 99999,
    },
]


def main():
    sys.stdout.reconfigure(encoding='utf-8')

    doc_path = Path('Part5_实战篇.docx')
    if not doc_path.exists():
        print(f"Error: {doc_path} not found")
        return

    doc = Document(str(doc_path))
    output_dir = Path('wechat_articles')
    output_dir.mkdir(exist_ok=True)

    print(f"文档共 {len(doc.paragraphs)} 个段落, {len(doc.tables)} 个表格\n")

    for idx, article in enumerate(ARTICLES, 1):
        print(f"[{idx}/{len(ARTICLES)}] {article['title']}")
        html = convert_chapter(
            doc, article['start'], article['end'],
            article['title'], article['subtitle'],
        )
        out_path = output_dir / article['file']
        out_path.write_text(html, encoding='utf-8')
        size_kb = out_path.stat().st_size / 1024
        print(f"     -> {out_path.name} ({size_kb:.0f} KB)\n")

    print(f"完成！共 {len(ARTICLES)} 篇，保存在 {output_dir}/")
    print("使用方法：浏览器打开 HTML → Ctrl+A 全选 → Ctrl+C 复制 → 粘贴到微信公众号编辑器")


if __name__ == '__main__':
    main()
