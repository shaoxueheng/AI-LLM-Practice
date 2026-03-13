"""
将 Part4_应用篇.docx 转换为多篇微信公众号 HTML 文章
复用 convert_to_wechat.py 的转换引擎
"""
import sys
from pathlib import Path
from docx import Document

# 复用 Part1 脚本的所有转换函数
from convert_to_wechat import convert_chapter

# ============================================================
# Part4 文章拆分方案（9篇）
# ============================================================
ARTICLES = [
    # --- 第13章 多模态大模型应用（拆为2篇） ---
    {
        "file": "24_多模态大模型全解析.html",
        "title": "大模型「五感觉醒」：多模态技术全解析",
        "subtitle": "VLM图像理解、Whisper语音识别、Stable Diffusion文生图、视频生成，一文看懂",
        "start": 2, "end": 205,
    },
    {
        "file": "25_多模态内容分析平台实战.html",
        "title": "实战！用Qwen2.5-VL+Whisper构建多模态内容分析平台",
        "subtitle": "图文音频一体化处理，附Gradio界面完整代码",
        "start": 205, "end": 400,
    },
    # --- 第14章 大模型应用开发实战（拆为2篇） ---
    {
        "file": "26_大模型应用架构与对话系统.html",
        "title": "大模型应用怎么设计？架构、FastAPI服务、对话系统全攻略",
        "subtitle": "三层架构设计、流式接口、Gradio/Streamlit界面、多轮对话意图识别",
        "start": 400, "end": 800,
    },
    {
        "file": "27_文档智能与智能写作助手实战.html",
        "title": "实战：合同审核+代码助手+端到端智能写作助手",
        "subtitle": "文档处理流水线、代码生成Prompt设计、FastAPI+前端完整工程代码",
        "start": 800, "end": 1176,
    },
    # --- 第15章 大模型部署与推理优化（拆为2篇） ---
    {
        "file": "28_大模型量化与推理框架.html",
        "title": "大模型推理太慢太贵？量化+推理框架帮你省钱提速",
        "subtitle": "INT4/INT8/GPTQ/AWQ量化对比、vLLM/TGI/Ollama框架选型、KV Cache优化、多GPU并行",
        "start": 1176, "end": 1351,
    },
    {
        "file": "29_边缘部署与vLLM高并发实战.html",
        "title": "手机上跑大模型？边缘部署+成本估算+vLLM生产级部署实战",
        "subtitle": "端侧模型推荐、GPU选型矩阵、云服务对比、Nginx负载均衡，附性能测试脚本",
        "start": 1351, "end": 1578,
    },
    # --- 第16章 大模型评测与安全（拆为3篇） ---
    {
        "file": "30_大模型评测体系全解析.html",
        "title": "大模型怎么评分？MMLU、C-Eval、HumanEval评测基准全解析",
        "subtitle": "通用与领域评测基准、LLM-as-Judge、自动化评测流水线搭建",
        "start": 1578, "end": 1864,
    },
    {
        "file": "31_大模型安全幻觉与红队测试.html",
        "title": "大模型会说谎、有偏见？幻觉、偏见、内容安全与红队测试详解",
        "subtitle": "安全风险分类、幻觉检测、内容过滤多层防护、自动化红队测试框架",
        "start": 1864, "end": 2253,
    },
    {
        "file": "32_评测安全审计系统实战.html",
        "title": "实战：搭建一套完整的大模型评测与安全审计系统",
        "subtitle": "统一配置管理、评测+安全双引擎、一键运行完整项目代码",
        "start": 2253, "end": 99999,
    },
]


def main():
    sys.stdout.reconfigure(encoding='utf-8')

    doc_path = Path('Part4_应用篇.docx')
    if not doc_path.exists():
        print(f"Error: {doc_path} not found")
        return

    doc = Document(str(doc_path))
    output_dir = Path('wechat_articles')
    output_dir.mkdir(exist_ok=True)

    print(f"文档共 {len(doc.paragraphs)} 个段落, {len(doc.tables)} 个表格\n")

    for idx, article in enumerate(ARTICLES, 1):
        print(f"[{idx}/{len(ARTICLES)}] {article['title']}")
        html = convert_chapter(
            doc, article['start'], article['end'],
            article['title'], article['subtitle'],
        )
        out_path = output_dir / article['file']
        out_path.write_text(html, encoding='utf-8')
        size_kb = out_path.stat().st_size / 1024
        print(f"     -> {out_path.name} ({size_kb:.0f} KB)\n")

    print(f"完成！共 {len(ARTICLES)} 篇，保存在 {output_dir}/")
    print("使用方法：浏览器打开 HTML → Ctrl+A 全选 → Ctrl+C 复制 → 粘贴到微信公众号编辑器")


if __name__ == '__main__':
    main()
